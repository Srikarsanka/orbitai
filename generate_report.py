"""
ORBIT AI - Final Report Document Generator
Reads the generated markdown documentation + existing front pages + annexure
and produces a properly formatted DOCX following SIST guidelines.

Font: Arial throughout
Spacing: 1.5 throughout
Page numbering: Roman (i,ii,iii) before Introduction, Arabic (1,2,3) from Introduction
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import os

# ============================================================
# CONFIGURATION
# ============================================================
OUTPUT_PATH = r'C:\ORBIT\ORBIT_AI_Final_Report.docx'
MD_PATH = r'C:\Users\sanka\.gemini\antigravity\brain\d837886b-b3e0-4ec3-b071-ad9da857c3a0\orbit_documentation.md'

FONT_NAME = 'Arial'
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_CHAPTER = Pt(14)
FONT_SIZE_SECTION = Pt(12)
FONT_SIZE_SUBSECTION = Pt(12)
LINE_SPACING = 1.5

MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.5)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_margins(section):
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

def make_run(para, text, bold=False, italic=False, size=FONT_SIZE_BODY, color=None, underline=False):
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rpr = run._element.get_or_add_rPr()
    rpr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'))
    if color:
        run.font.color.rgb = color
    return run

def add_para(doc, text, bold=False, italic=False, size=FONT_SIZE_BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), space_before=Pt(0)):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = space_after
    pf.space_before = space_before
    make_run(para, text, bold=bold, italic=italic, size=size)
    return para

def add_chapter_heading(doc, text):
    """Chapter heading: 14pt, Bold, Centered, ALL CAPS"""
    para = add_para(doc, text.upper(), bold=True, size=FONT_SIZE_CHAPTER, 
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(24), space_after=Pt(12))
    return para

def add_section_heading(doc, text):
    """Section heading: 12pt, Bold"""
    para = add_para(doc, text, bold=True, size=FONT_SIZE_SECTION,
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(18), space_after=Pt(6))
    return para

def add_subsection_heading(doc, text):
    """Subsection heading: 12pt, Bold, Italic"""
    para = add_para(doc, text, bold=True, italic=True, size=FONT_SIZE_SUBSECTION,
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(6))
    return para

def add_body_text(doc, text):
    """Body text: 12pt, Justified"""
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

def add_bullet(doc, text, bold_prefix=""):
    """Bulleted body text"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(3)
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-0.5)
    if bold_prefix:
        make_run(para, "• " + bold_prefix, bold=True)
        make_run(para, text)
    else:
        make_run(para, "• " + text)
    return para

def add_figure_placeholder(doc, fig_num, caption, prompt):
    """Add figure placeholder with blank space and caption"""
    # Blank space for figure
    para = add_para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(24))
    add_para(doc, f"[Insert {fig_num} here]", bold=True, italic=True, 
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_para(doc, f"Prompt: {prompt}", italic=True, size=Pt(10),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    # Add some blank lines for space
    for _ in range(4):
        add_para(doc, "", space_after=Pt(12))
    # Caption
    add_para(doc, f"({fig_num}) {caption}", bold=True, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def add_roman_page_numbers(section):
    """Add Roman numeral page numbering to a section"""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE  \\* ROMAN </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run = para.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

def add_arabic_page_numbers(section, restart=True):
    """Add Arabic numeral page numbering to a section, optionally restarting from 1"""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Restart page numbering
    if restart:
        sectPr = section._sectPr
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="1"/>')
        sectPr.append(pgNumType)
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run = para.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

import docx

# ============================================================
# READ MARKDOWN CONTENT
# ============================================================
print("Reading markdown documentation...")
with open(MD_PATH, 'r', encoding='utf-8') as f:
    md_content = f.read()

# ============================================================
# CREATE DOCUMENT
# ============================================================
print("Creating document...")
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = FONT_SIZE_BODY
style.paragraph_format.line_spacing = LINE_SPACING

# Set default margins
section = doc.sections[0]
set_margins(section)

# ============================================================
# TITLE PAGE (No Page Number)
# ============================================================
print("  Adding title page...")
# Add blank lines for spacing
for _ in range(3):
    add_para(doc, "")

add_para(doc, "SATHYABAMA", bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "INSTITUTE OF SCIENCE AND TECHNOLOGY", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "(DEEMED TO BE UNIVERSITY)", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Accredited with Grade \"A++\" by NAAC  |  Approved by AICTE", size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "JEPPIAAR NAGAR, RAJIV GANDHI SALAI, CHENNAI - 600 119", size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

add_para(doc, "")
add_para(doc, "AI POWERED VIRTUAL CLASSROOM FOR PERSONALIZED LEARNING", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
add_para(doc, "")

add_para(doc, "Submitted in partial fulfillment of the requirements for the award of", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Bachelor of Technology", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "in", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "COMPUTER SCIENCE AND ENGINEERING", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))

add_para(doc, "")
add_para(doc, "By", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "SANKA SARAN SAI KRISHNA SRIKAR (Reg No. 42111148)", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "SANDADI MOHAN KANTHU KUMAR REDDY", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "")

add_para(doc, "Under the Guidance of", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Dr. M SELVI, Associate Professor", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "")

add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "SCHOOL OF COMPUTING", bold=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "")
add_para(doc, "APRIL – 2026", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ============================================================
# BONAFIDE CERTIFICATE (No Page Number)
# ============================================================
print("  Adding bonafide certificate...")
add_para(doc, "")
add_para(doc, "BONAFIDE CERTIFICATE", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para(doc, "")

add_body_text(doc, 
    "This is to certify that this project report titled \"AI POWERED VIRTUAL CLASSROOM FOR PERSONALIZED LEARNING\" "
    "is the bonafide work of SANKA SARAN SAI KRISHNA SRIKAR (42111148) and SANDADI MOHAN KANTHU KUMAR REDDY, "
    "who carried out the project work under my supervision. Certified further that to the best of our knowledge "
    "the work reported herein does not form part of any other thesis or dissertation on the basis of which a degree "
    "or award was conferred on an earlier occasion on this or any other candidate.")

add_para(doc, "")
add_para(doc, "")
add_para(doc, "")

# Signature blocks
add_para(doc, "Dr. M SELVI", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Project Guide", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Associate Professor", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Department of CSE", align=WD_ALIGN_PARAGRAPH.LEFT)

add_para(doc, "")
add_para(doc, "")

add_para(doc, "Head of the Department", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Department of Computer Science and Engineering", align=WD_ALIGN_PARAGRAPH.LEFT)

add_para(doc, "")
add_para(doc, "Internal Examiner                                                    External Examiner", 
         align=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break(doc)

# ============================================================
# DECLARATION (No Page Number)
# ============================================================
print("  Adding declaration...")
add_para(doc, "DECLARATION", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para(doc, "")

add_body_text(doc,
    "I hereby declare that the project report titled \"AI POWERED VIRTUAL CLASSROOM FOR PERSONALIZED LEARNING\" "
    "submitted to Sathyabama Institute of Science and Technology in partial fulfillment of the requirements for "
    "the award of Bachelor of Technology in Computer Science and Engineering is a record of original work done by "
    "me under the guidance of Dr. M Selvi, Associate Professor, Department of Computer Science and Engineering, "
    "Sathyabama Institute of Science and Technology, Chennai – 600 119.")

add_para(doc, "")
add_body_text(doc, "I further declare that the work reported in this project has not been submitted and will not be "
    "submitted, either in part or in full, for the award of any other degree or diploma in this institute or "
    "any other institute or university.")
    
add_para(doc, "")
add_para(doc, "")
add_para(doc, "")
add_para(doc, "Place: Chennai", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "Date:", align=WD_ALIGN_PARAGRAPH.LEFT)
add_para(doc, "")
add_para(doc, "(SANKA SARAN SAI KRISHNA SRIKAR)                    (SANDADI MOHAN KANTHU KUMAR REDDY)", 
         bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break(doc)

# ============================================================
# ACKNOWLEDGEMENT (Roman numeral page i)
# ============================================================
print("  Adding acknowledgement...")
# New section for roman page numbers
new_section = doc.add_section()
set_margins(new_section)
add_roman_page_numbers(new_section)

add_para(doc, "ACKNOWLEDGEMENT", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para(doc, "")

add_body_text(doc, 
    "I am grateful to the Chancellor, Vice Chancellor, Director (Academic Administration), and Registrar "
    "of Sathyabama Institute of Science and Technology, Chennai, for providing the necessary infrastructure "
    "and resources for the successful completion of this project.")
add_body_text(doc,
    "I express my sincere thanks to Dr. S. Vigneshwari, M.E., Ph.D., Head of the Department of Computer Science "
    "and Engineering, for her support and encouragement throughout this project.")
add_body_text(doc,
    "I would like to express my sincere and deep sense of gratitude to my Project Guide Dr. M Selvi, "
    "Associate Professor, Department of Computer Science and Engineering, for her valuable guidance, "
    "constant encouragement, and constructive criticism throughout the course of this project work.")
add_body_text(doc,
    "I wish to express my thanks to all Teaching and Non-teaching staff members of the Department of "
    "Computer Science and Engineering who were helpful in many ways for the completion of the project.")
add_body_text(doc,
    "Finally, I thank my parents, family, and friends for their continuous encouragement, moral support, "
    "and motivation throughout the course of this project.")

add_page_break(doc)

# ============================================================
# ABSTRACT (Roman numeral page)
# ============================================================
print("  Adding abstract...")
add_para(doc, "ABSTRACT", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
add_para(doc, "")

add_body_text(doc,
    "The sudden shift to web-based learning has resulted in the extensive adoption of virtual classroom platforms. "
    "However, most existing tools are designed for generic video conferencing and lack education-specific features "
    "such as biometric attendance verification, integrated code execution, and multilingual content accessibility. "
    "This project presents ORBIT AI, an AI-powered virtual classroom platform built on a microservices architecture "
    "using Angular 18, Node.js (Express.js), Python (FastAPI), MongoDB, and the Agora WebRTC SDK.")
add_body_text(doc,
    "The platform implements biometric face-recognition-based authentication and automated attendance tracking "
    "using InsightFace's buffalo_l deep learning model, which generates 512-dimensional facial embeddings for "
    "cosine similarity matching at configurable intervals during live sessions. A sandboxed, multi-language code "
    "execution environment powered by the Monaco Editor supports Python, C, C++, and Java, augmented by AI-assisted "
    "error correction. A collaborative whiteboard with Socket.io-based real-time synchronization enables visual "
    "instruction, while a Dockerized voice translation microservice (FFmpeg, OpenAI Whisper, gTTS) provides "
    "multilingual accessibility for recorded lectures.")
add_body_text(doc,
    "The system is deployed on a cloud-native infrastructure using Render.com, MongoDB Atlas, and Cloudinary, "
    "demonstrating scalability and cost-efficiency. Experimental results confirm reliable biometric attendance "
    "(2.1s average latency per capture), stable multi-participant video conferencing, and accurate code execution "
    "across supported languages. ORBIT AI represents a comprehensive, purpose-built solution for modern "
    "online education.")
add_para(doc, "")
add_body_text(doc, "Keywords: Virtual Classroom, Face Recognition, InsightFace, WebRTC, Biometric Attendance, "
    "Code Execution, Voice Translation, Microservices, Angular, Node.js, FastAPI, MongoDB.")

add_page_break(doc)

# ============================================================
# TABLE OF CONTENTS (Roman numeral page)
# ============================================================
print("  Adding table of contents...")
add_para(doc, "TABLE OF CONTENTS", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

toc_entries = [
    ("", "ACKNOWLEDGEMENT", "i"),
    ("", "ABSTRACT", "ii"),
    ("", "TABLE OF CONTENTS", "iii"),
    ("", "LIST OF FIGURES", "v"),
    ("", "", ""),
    ("CHAPTER 1", "INTRODUCTION", "1"),
    ("1.1", "Background & Motivation", "1"),
    ("1.2", "Problem Statement", "3"),
    ("", "", ""),
    ("CHAPTER 2", "LITERATURE SURVEY", "5"),
    ("2.1", "Review of Existing Systems", "5"),
    ("2.2", "Inferences from Literature Survey", "14"),
    ("", "", ""),
    ("CHAPTER 3", "AIM AND SCOPE OF THE PRESENT INVESTIGATION", "17"),
    ("3.1", "Aim of the Project", "17"),
    ("3.2", "Scope of the Present Investigation", "18"),
    ("3.3", "Feasibility Studies / Risk Analysis", "19"),
    ("3.4", "Software Requirements Specification", "21"),
    ("", "", ""),
    ("CHAPTER 4", "MATERIALS AND METHODS; ALGORITHMS USED", "25"),
    ("4.1", "Selected Methodologies", "25"),
    ("4.2", "Architecture Diagram", "26"),
    ("4.3", "Detailed Description of Modules and Workflow", "29"),
    ("4.4", "Estimated Cost for Implementation", "48"),
    ("4.5", "Algorithms and Techniques Used", "49"),
    ("", "", ""),
    ("CHAPTER 5", "RESULTS AND DISCUSSION, PERFORMANCE ANALYSIS", "53"),
    ("5.1", "Achieved Outcomes and System Performance", "53"),
    ("", "", ""),
    ("CHAPTER 6", "SUMMARY AND CONCLUSIONS", "60"),
    ("6.1", "Summary and Key Findings", "60"),
    ("6.2", "Future Scope and Improvements", "61"),
    ("", "", ""),
    ("", "REFERENCES", "63"),
    ("", "ANNEXURES", "66"),
]

for num, title, page in toc_entries:
    if not title and not num:
        add_para(doc, "", space_after=Pt(2))
        continue
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(3)
    pf.tab_stops.add_tab_stop(Cm(14.5), alignment=docx.enum.text.WD_TAB_ALIGNMENT.RIGHT, leader=docx.enum.text.WD_TAB_LEADER.DOTS)
    
    is_chapter = num.startswith("CHAPTER")
    if is_chapter:
        make_run(para, f"{num}: {title}", bold=True)
    elif num:
        make_run(para, f"    {num}  {title}")
    else:
        make_run(para, f"{title}", bold=True)
    make_run(para, f"\t{page}")

add_page_break(doc)

# ============================================================
# LIST OF FIGURES (Roman numeral page)
# ============================================================
print("  Adding list of figures...")
add_para(doc, "LIST OF FIGURES", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

figures = [
    ("Fig. 1.1", "Global E-Learning Market Growth (2010–2025)", "2"),
    ("Fig. 4.1", "ORBIT AI — Overall System Architecture", "27"),
    ("Fig. 4.2", "ORBIT AI — MongoDB Schema Entity Relationships", "28"),
    ("Fig. 4.3", "Authentication & Registration Flow", "31"),
    ("Fig. 4.4", "Class Management Workflow", "33"),
    ("Fig. 4.5", "Class Session Lifecycle", "36"),
    ("Fig. 4.6", "Biometric Attendance Workflow", "39"),
    ("Fig. 4.7", "Materials Management Module Flow", "41"),
    ("Fig. 4.8", "Code Execution & AI Assistance Flow", "43"),
    ("Fig. 4.9", "Collaborative Whiteboard Real-Time Architecture", "45"),
    ("Fig. 4.10", "Voice Translation Pipeline", "46"),
    ("Fig. 4.11", "Announcements & News Feed Module", "47"),
    ("Fig. 4.12", "Analytics Dashboard Architecture", "48"),
    ("Fig. 5.1", "Login Page and Biometric Authentication Screen", "54"),
    ("Fig. 5.2", "Student and Faculty Dashboard Views", "55"),
    ("Fig. 5.3", "Live Video Session with Multiple Participants", "55"),
    ("Fig. 5.4", "Session Attendance Report View", "56"),
    ("Fig. 5.5", "Integrated Code Execution Environment", "57"),
    ("Fig. 5.6", "Collaborative Whiteboard in Live Session", "57"),
    ("Fig. 5.7", "Materials Library (Student View)", "58"),
    ("Fig. 5.8", "Voice-Translated Lecture Recording Player", "58"),
    ("Fig. 5.9", "Faculty Analytics Dashboard", "59"),
]

for fig_num, caption, page in figures:
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(3)
    pf.tab_stops.add_tab_stop(Cm(14.5), alignment=docx.enum.text.WD_TAB_ALIGNMENT.RIGHT, leader=docx.enum.text.WD_TAB_LEADER.DOTS)
    make_run(para, f"{fig_num}  {caption}")
    make_run(para, f"\t{page}")

add_page_break(doc)

# ============================================================
# PARSE AND ADD MARKDOWN CONTENT
# ============================================================
print("  Parsing markdown content...")

# Start new section for Arabic numbering
new_section = doc.add_section()
set_margins(new_section)
add_arabic_page_numbers(new_section, restart=True)

lines = md_content.split('\n')
i = 0
skip_header = True  # Skip the metadata header block

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # Skip the metadata header (title, tech stack, etc.)
    if skip_header:
        if stripped.startswith('## Chapter'):
            skip_header = False
        else:
            i += 1
            continue
    
    # Skip empty lines and horizontal rules
    if not stripped or stripped == '---':
        i += 1
        continue
    
    # Chapter headings: ## Chapter – X: TITLE
    if stripped.startswith('## Chapter'):
        title = stripped.replace('## ', '').strip()
        add_chapter_heading(doc, title)
        i += 1
        continue
    
    # Section headings: ### X.X Title
    if stripped.startswith('### '):
        title = stripped.replace('### ', '').strip()
        add_section_heading(doc, title)
        i += 1
        continue
    
    # Subsection headings: #### X.X.X Title or #### Module X:
    if stripped.startswith('#### '):
        title = stripped.replace('#### ', '').strip()
        add_subsection_heading(doc, title)
        i += 1
        continue
    
    # Figure placeholders: > **[INSERT FIGURE HERE]** or > **[INSERT SCREENSHOT HERE]**
    if stripped.startswith('> **[INSERT'):
        # Collect the prompt and caption
        prompt_text = ""
        caption_text = ""
        fig_num = ""
        j = i + 1
        while j < len(lines):
            sline = lines[j].strip()
            if sline.startswith('> *Diagram Generation Prompt') or sline.startswith('> *Screenshot:'):
                prompt_text = sline.replace('> ', '').strip('*').strip()
            elif sline.startswith('> "') or sline.startswith('> \''):
                prompt_text += " " + sline.replace('> ', '').strip('"\'')
            elif sline.startswith('*(Fig.') or sline.startswith('*(Fig '):
                caption_text = sline.strip('*()').strip()
                fig_match = re.search(r'Fig\.?\s*(\d+\.\d+)', sline)
                if fig_match:
                    fig_num = f"Fig. {fig_match.group(1)}"
                j += 1
                break
            elif not sline or sline == '---':
                j += 1
                break
            j += 1
        
        if fig_num:
            add_figure_placeholder(doc, fig_num, caption_text, prompt_text[:300] if prompt_text else "Generate appropriate diagram")
        i = j
        continue
    
    # Blockquotes (keeping Note to Author etc.)
    if stripped.startswith('> '):
        text = stripped[2:].strip('*').strip()
        if text and not text.startswith('[INSERT'):
            add_para(doc, text, italic=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.LEFT)
        i += 1
        continue
    
    # Bold paragraph headers: **Header:** Text
    bold_match = re.match(r'\*\*(.+?)\*\*(.+)?', stripped)
    if bold_match and not stripped.startswith('- '):
        bold_text = bold_match.group(1)
        rest = bold_match.group(2) if bold_match.group(2) else ""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_after = Pt(6)
        make_run(para, bold_text, bold=True)
        if rest:
            make_run(para, rest.strip())
        i += 1
        continue
    
    # Bullet points
    if stripped.startswith('- '):
        text = stripped[2:].strip()
        # Check for bold prefix: - **text:** rest
        bm = re.match(r'\*\*(.+?)\*\*\s*(.*)', text)
        if bm:
            add_bullet(doc, bm.group(2), bold_prefix=bm.group(1) + " ")
        else:
            # Remove markdown formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            add_bullet(doc, text)
        i += 1
        continue
    
    # Code blocks
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            code_text = '\n'.join(code_lines)
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.line_spacing = 1.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.left_indent = Cm(1.0)
            make_run(para, code_text, size=Pt(10))
        i += 1  # skip closing ```
        continue
    
    # Table: | ... | ... |
    if stripped.startswith('|'):
        # Collect all table rows
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row_line = lines[i].strip()
            if not row_line.startswith('|---') and not row_line.startswith('| ---'):
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                table_rows.append(cells)
            i += 1
        
        if table_rows:
            num_cols = max(len(row) for row in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = 'Table Grid'
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < num_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        pf = para.paragraph_format
                        pf.line_spacing = LINE_SPACING
                        is_header = (ri == 0)
                        cell_text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                        make_run(para, cell_text_clean, bold=is_header, size=Pt(10))
            add_para(doc, "")
        continue
    
    # Regular paragraph — clean up markdown formatting
    text = stripped
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # italic
    text = re.sub(r'`(.+?)`', r'\1', text)  # code
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
    
    if text and len(text) > 3:
        add_body_text(doc, text)
    
    i += 1

# ============================================================
# ANNEXURE SECTION
# ============================================================
print("  Adding annexures...")
add_page_break(doc)
add_chapter_heading(doc, "ANNEXURES")
add_para(doc, "")

# Annexure 1: Plagiarism Report Placeholder
add_section_heading(doc, "ANNEXURE 1")
add_body_text(doc, "Plagiarism Report (Report taken in Turnitin or iThenticate and not exceeding 25%)")
add_para(doc, "")
add_para(doc, "[Insert Plagiarism Report Here]", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    add_para(doc, "")

# Annexure 2: Research Paper
add_page_break(doc)
add_section_heading(doc, "ANNEXURE 2")
add_body_text(doc, "Published Paper / Conference Paper / Journal Paper / Patent Published (if Applicable)")
add_para(doc, "")
add_body_text(doc, "Title: AI-Powered Virtual Classroom for Personalised Learning")
add_body_text(doc, "Authors: Sanka Saran Sai Krishna Srikar, Sandadi Mohan Kanthu Kumar Reddy, Dr. M Selvi, "
    "Dr. G Kalaiarasi, Yogitha R, Dr. Joshila Grace L K")
add_body_text(doc, "Institution: Sathyabama Institute of Science and Technology, Chennai, India")
add_para(doc, "")
add_para(doc, "[Insert Research Paper Here]", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    add_para(doc, "")

# Annexure 3: Project Screenshots
add_page_break(doc)
add_section_heading(doc, "ANNEXURE 3")
add_body_text(doc, "Additional Photographs / Screenshots taken during the project (which are not used inside the thesis).")
add_para(doc, "")
add_para(doc, "[Insert Additional Project Screenshots Here]", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    add_para(doc, "")

# Annexure 4: Individual Photo
add_page_break(doc)
add_section_heading(doc, "ANNEXURE 4")
add_body_text(doc, "Individual Photograph of the Candidate(s)")
add_para(doc, "")
add_para(doc, "[Insert Individual Photographs Here]", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    add_para(doc, "")

# Annexure 5: SDG Mapping
add_page_break(doc)
add_section_heading(doc, "ANNEXURE 5")
add_body_text(doc, "SDG MAPPING")
add_para(doc, "")
add_body_text(doc, "AI POWERED VIRTUAL CLASSROOM FOR PERSONALIZED LEARNING")
add_para(doc, "")

# SDG Goal Table
sdg_table = doc.add_table(rows=4, cols=3)
sdg_table.style = 'Table Grid'
sdg_headers = ['SDG Goal', 'Description', 'Mapped Program Outcomes']
for ci, h in enumerate(sdg_headers):
    cell = sdg_table.cell(0, ci)
    cell.text = ""
    para = cell.paragraphs[0]
    make_run(para, h, bold=True, size=Pt(11))

sdg_data = [
    ('SDG 4', 'Quality Education', 'PO1, PO3, PO5, PO6, PO12'),
    ('SDG 9', 'Industry, Innovation and Infrastructure', 'PO1, PO2, PO3, PO5'),
    ('SDG 10', 'Reduced Inequalities', 'PO6, PO7, PO12'),
]
for ri, (goal, desc, pos) in enumerate(sdg_data):
    sdg_table.cell(ri+1, 0).text = goal
    sdg_table.cell(ri+1, 1).text = desc
    sdg_table.cell(ri+1, 2).text = pos

add_para(doc, "")

# PO-PSO Mapping Table
add_body_text(doc, "PO-PSO Mapping:")
add_para(doc, "")

po_table = doc.add_table(rows=16, cols=3)
po_table.style = 'Table Grid'
po_headers = ['PO-PSO', 'Mapping Level', 'Justification']
for ci, h in enumerate(po_headers):
    cell = po_table.cell(0, ci)
    cell.text = ""
    para = cell.paragraphs[0]
    make_run(para, h, bold=True, size=Pt(11))

po_data = [
    ('PO1', '3', 'Application of AI/ML algorithms for face recognition and code execution'),
    ('PO2', '3', 'Analysis of biometric attendance and engagement tracking challenges'),
    ('PO3', '3', 'Design of intelligent virtual classroom system'),
    ('PO4', '2', 'Experimental validation using live sessions and test data'),
    ('PO5', '3', 'Use of Angular, Node.js, Python, MongoDB, Docker, Agora SDK'),
    ('PO6', '3', 'Strong societal relevance — accessible, inclusive education'),
    ('PO7', '2', 'Multilingual translation promotes language inclusivity'),
    ('PO8', '1', 'Ethical responsibility in biometric data handling'),
    ('PO9', '2', 'Team-based project execution'),
    ('PO10', '2', 'Documentation and academic presentation'),
    ('PO11', '1', 'Planning and execution of development sprints'),
    ('PO12', '3', 'Learning emerging AI/ML and cloud-native techniques'),
    ('PSO1', '3', 'Computational modeling for face recognition and attendance'),
    ('PSO2', '3', 'AI/ML-based intelligent educational system'),
    ('PSO3', '3', 'Development of real-world cloud-deployed platform'),
]
for ri, (po, level, just) in enumerate(po_data):
    po_table.cell(ri+1, 0).text = po
    po_table.cell(ri+1, 1).text = level
    po_table.cell(ri+1, 2).text = just

# ============================================================
# SAVE
# ============================================================
print(f"\nSaving document to: {OUTPUT_PATH}")
doc.save(OUTPUT_PATH)
print("DONE! Document generated successfully.")
print(f"File: {OUTPUT_PATH}")
